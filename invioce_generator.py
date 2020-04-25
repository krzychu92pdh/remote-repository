from docx import Document
import datetime
import config
import pln_generator as pln
from docx.enum.table import WD_TABLE_ALIGNMENT
import subprocess
import smtplib
from email.message import EmailMessage
import webbrowser

document = Document('template.docx')


def date_invoice():
	#ostatni dzień poprzedniego miesiąca dla faktury miesięcznej
	today = datetime.date.today()
	first = today.replace(day=1)
	lastday = first - datetime.timedelta(days=1)
	lastdayofmonth = lastday.strftime("%d/%m/%Y")
	return lastdayofmonth

def monthandyear_invoice():
	today = datetime.date.today()
	first = today.replace(day=1)
	lastday = first - datetime.timedelta(days=1)
	monthandyear = lastday.strftime("%m/%Y")
	return monthandyear

def fourteen_days():
	today = datetime.date.today()
	first = today.replace(day=1)
	lastday = first + datetime.timedelta(days=13)
	termin = lastday.strftime("%d/%m/%Y")
	return termin

def findandinsert_par(find, insert):
    for paragraph in document.paragraphs:
        if find in paragraph.text:
            paragraph.add_run(insert).bold = True


def add_row():
	lp = 0
	suma = []
	while True:
		lp += 1
		x = input('Wpisz "Y" jesli chcessz wprowadzić usługę do faktury lub "N" żeby skonczyć: ')
		print(" ")
		if x == 'y':
			nazwa = input('Podaj nazwę usługi: ')
			netto = input('Podaj cenne netto usługi: ')
			netto_int = int(netto)
			suma.append(netto_int)


			document.tables[1].add_row().cells
			document.tables[1].cell(-1,0).text = str(lp)
			document.tables[1].cell(-1,1).text = nazwa
			document.tables[1].cell(-1,3).text = 'szt.'
			document.tables[1].cell(-1,4).text = '1'
			document.tables[1].cell(-1,4).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
			document.tables[1].cell(-1,5).text = netto
			document.tables[1].cell(-1,5).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
			document.tables[1].cell(-1,6).text = 'zw'
			document.tables[1].cell(-1,6).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
			document.tables[1].cell(-1,7).text = netto
			document.tables[1].cell(-1,7).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
			document.tables[1].cell(-1,8).text = "0"
			document.tables[1].cell(-1,8).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
			document.tables[1].cell(-1,9).text = netto
			document.tables[1].cell(-1,9).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
		elif x == 'n':
			break
	global pelna_kwota
	pelna_kwota = str(sum(suma))


def doc2pdf_linux(doc):
    cmd = 'libreoffice --convert-to pdf'.split() + [doc]
    p = subprocess.Popen(cmd, stderr=subprocess.PIPE, stdout=subprocess.PIPE)
    stdout, stderr = p.communicate()


def sendEmail(x):
    EMAIL_ADRESS= config.email
    EMAIL_PASSWORD= config.tokenpass

    msg = EmailMessage()
    msg['Subject'] = config.subject
    msg['From'] = EMAIL_ADRESS
    msg['To'] = config.emailto
    msg.set_content(config.content)

    files = []
    files.append(x)

    for file in files:
        with open(file, 'rb') as f:
            file_data = f.read()
            file_name = "invoice.pdf"

    msg.add_attachment(file_data, maintype='application', subtype='octet=stream', filename=file_name)

    with smtplib.SMTP_SSL('smtp.gmail.com', 465) as smtp:

        smtp.login(EMAIL_ADRESS, EMAIL_PASSWORD)

        smtp.send_message(msg)



#### DOCUMENT ####



#nagłówek
number_invoice = input("Wpisz numer faktury:")
findandinsert_par("Data wystawienia:",date_invoice())
findandinsert_par("Faktura nr FV", str(number_invoice) + "/" + monthandyear_invoice()) 
findandinsert_par("Miejsce wystawienia:",config.city)

#dane sprzedawcy:
document.tables[0].cell(1,0).paragraphs[0].add_run(config.name).bold = True
document.tables[0].cell(2,0).text = config.adress1
document.tables[0].cell(3,0).text = config.adress2
document.tables[0].cell(4,0).text = config.nip
document.tables[0].cell(5,0).text = config.email
document.tables[0].cell(6,0).text = config.phone

#dane nabywcy:
document.tables[0].cell(1,1).paragraphs[0].add_run(config.name_n).bold = True
document.tables[0].cell(2,1).text = config.adress1_n
document.tables[0].cell(3,1).text = config.adress2_n
document.tables[0].cell(4,1).text = config.nip_n


#tabela z usługami
add_row()

#tabela zbiorcza
document.tables[2].cell(1,1).text = pelna_kwota
document.tables[2].cell(1,1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
document.tables[2].cell(2,1).text = pelna_kwota
document.tables[2].cell(2,1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
document.tables[2].cell(1,3).text = pelna_kwota
document.tables[2].cell(1,3).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
document.tables[2].cell(2,3).text = pelna_kwota
document.tables[2].cell(2,3).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

#tabela z informacjami
document.tables[3].cell(2,1).paragraphs[0].add_run(config.bank).bold = True
document.tables[3].cell(3,1).text = config.account_nr
document.tables[3].cell(1,1).text = fourteen_days()
document.tables[3].cell(3,3).text = pln.generator(pelna_kwota)
document.tables[3].cell(0,3).text = pelna_kwota + " zł"
document.tables[3].cell(2,3).text = pelna_kwota + " zł"

#osoba upoważniona do odbioru
document.tables[4].cell(0,0).paragraphs[0].add_run(config.name).bold = True


print(pelna_kwota)

document.save('invoice.docx')

#zapis do pdf
doc2pdf_linux('invoice.docx')

#podgląd dokumentu
webbrowser.open_new(r"invoice.pdf")


#wysłanie mailem
zgoda = input("Jeśli chcesz wysłać fakturkę mailem napisz 'Y' jeśli nie to 'N': ")
if zgoda == "y":
    sendEmail("invoice.pdf")
    print("gotowe :)")
else:
    print("Mail nie został wysłany popraw fakturę w pliku docx")